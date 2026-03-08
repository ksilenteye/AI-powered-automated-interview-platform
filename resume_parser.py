"""
Enhanced Resume Parser with OCR and AI-based Categorization
Uses Tesseract OCR for scanned PDFs and Groq for skill categorization
Author: Kavya Bhardwaj
Features:
- Robust text extraction from PDFs, DOCX, and images
- Heuristic detection of scanned documents with OCR fallback
- Comprehensive skill extraction using regex and keyword matching
- AI-powered skill categorization into technical, tools, domain, and soft skills
- Graceful error handling and fallback mechanisms
"""

import re
import cv2
import numpy as np
import pytesseract
from PIL import Image
import PyPDF2
from pdf2image import convert_from_path
from docx import Document
from groq import Groq
import json


class EnhancedResumeParser:
    def __init__(self, groq_api_key=None):
        self.groq_client = Groq(api_key=groq_api_key) if groq_api_key else None

        # Comprehensive skills database
        self.skill_categories = {
            'programming_languages': [
                'python', 'javascript', 'java', 'c++', 'c#', 'ruby', 'php', 
                'swift', 'kotlin', 'go', 'rust', 'typescript', 'scala', 'r',
                'matlab', 'perl', 'dart', 'objective-c'
            ],
            'web_frameworks': [
                'react', 'angular', 'vue', 'django', 'flask', 'fastapi',
                'express', 'node.js', 'spring', 'asp.net', 'laravel',
                'next.js', 'nuxt.js', 'svelte', 'ember'
            ],
            'databases': [
                'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle',
                'cassandra', 'dynamodb', 'elasticsearch', 'neo4j', 'sqlite',
                'mariadb', 'couchdb'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean',
                'ibm cloud', 'oracle cloud', 'alibaba cloud'
            ],
            'devops_tools': [
                'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab',
                'terraform', 'ansible', 'puppet', 'chef', 'circleci', 'travis ci',
                'ci/cd', 'bitbucket'
            ],
            'ml_ai': [
                'machine learning', 'deep learning', 'tensorflow', 'pytorch',
                'keras', 'scikit-learn', 'pandas', 'numpy', 'opencv', 'nltk',
                'spacy', 'hugging face', 'computer vision', 'nlp', 'neural networks',
                'transformers', 'gan', 'cnn', 'rnn', 'lstm'
            ],
            'data_tools': [
                'tableau', 'power bi', 'excel', 'jupyter', 'spark', 'hadoop',
                'airflow', 'kafka', 'snowflake', 'databricks'
            ],
            'mobile': [
                'android', 'ios', 'react native', 'flutter', 'xamarin',
                'ionic', 'cordova'
            ],
            'soft_skills': [
                'leadership', 'communication', 'teamwork', 'problem solving',
                'analytical', 'critical thinking', 'time management', 'agile',
                'scrum', 'collaboration', 'presentation'
            ]
        }

    def parse_file(self, file_path):
        """Main entry point for parsing resume"""
        try:
            # Extract text based on file type
            if file_path.endswith('.pdf'):
                text = self._extract_pdf_with_fallback(file_path)
            elif file_path.endswith('.docx'):
                text = self._extract_docx(file_path)
            elif file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                text = self._extract_image_ocr(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()

            # Extract structured information
            resume_data = self._extract_info(text)

            # Categorize skills using AI if available
            if self.groq_client and resume_data['skills']:
                resume_data['skills_by_category'] = self._categorize_skills_with_ai(
                    resume_data['skills'], text
                )
            else:
                resume_data['skills_by_category'] = self._categorize_skills_manual(
                    resume_data['skills']
                )

            resume_data['raw_text'] = text
            return resume_data

        except Exception as e:
            return {
                'error': str(e),
                'name': 'Unknown',
                'email': '',
                'phone': '',
                'skills': [],
                'skills_by_category': {},
                'experience': [],
                'education': [],
                'location': '',
                'raw_text': ''
            }

    def _extract_pdf_with_fallback(self, file_path):
        """Extract text from PDF with OCR fallback for scanned documents"""
        text = ""

        try:
            # Try regular PDF text extraction first
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            # Check if extraction was successful (not a scanned PDF)
            if len(text.strip()) > 100 and not self._is_likely_scanned(text):
                return text

            # Fallback to OCR for scanned PDFs
            print("Detected scanned PDF, using OCR...")
            return self._extract_pdf_with_ocr(file_path)

        except Exception as e:
            print(f"PDF extraction error: {e}, trying OCR...")
            return self._extract_pdf_with_ocr(file_path)

    def _is_likely_scanned(self, text):
        """Heuristic to detect if PDF is scanned (poor text extraction)"""
        if len(text.strip()) < 100:
            return True

        # Check for excessive whitespace or garbage characters
        words = text.split()
        if len(words) < 50:
            return True

        # Check for abnormal character distribution
        alpha_ratio = sum(c.isalpha() for c in text) / max(len(text), 1)
        if alpha_ratio < 0.5:
            return True

        return False

    def _extract_pdf_with_ocr(self, file_path):
        """Convert PDF pages to images and apply OCR"""
        try:
            # Convert PDF to images
            images = convert_from_path(
                file_path,
                dpi=300,
                poppler_path=r"C:\Users\asus gaming\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
            )

            text = ""
            for i, image in enumerate(images):
                # Convert PIL Image to numpy array for OpenCV
                img_array = np.array(image)

                # Preprocess image for better OCR
                processed = self._preprocess_image_for_ocr(img_array)

                # Apply OCR
                page_text = pytesseract.image_to_string(processed, config='--psm 6')
                text += page_text + "\n"

            return text

        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""

    def _extract_image_ocr(self, file_path):
        """Extract text from image file using OCR"""
        try:
            # Read image
            image = cv2.imread(file_path)

            # Preprocess
            processed = self._preprocess_image_for_ocr(image)

            # OCR
            text = pytesseract.image_to_string(processed, config='--psm 6')
            return text

        except Exception as e:
            print(f"Image OCR failed: {e}")
            return ""

    def _preprocess_image_for_ocr(self, image):
        """Preprocess image to improve OCR accuracy"""
        # Convert to grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image

        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)

        # Thresholding
        _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Deskew (optional)
        # Could add rotation correction here if needed

        return binary

    def _extract_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text

            return text
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""

    def _extract_info(self, text):
        """Extract structured information from text"""
        return {
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'skills': self._extract_skills(text),
            'experience': self._extract_experience(text),
            'education': self._extract_education(text),
            'location': self._extract_location(text),
            'projects': self._extract_projects(text)
        }

    def _extract_name(self, text):
        """Extract candidate name"""
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            # Usually name is in first few lines
            for line in lines[:5]:
                # Skip email and phone lines
                if '@' not in line and not re.search(r'\d{10}', line):
                    # Check if line looks like a name (2-4 words, capitalized)
                    words = line.split()
                    if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
                        return line
            return lines[0]
        return 'Unknown'

    def _extract_email(self, text):
        """Extract email address"""
        pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(pattern, text)
        return match.group(0) if match else ''

    def _extract_phone(self, text):
        """Extract phone number"""
        patterns = [
            r'\+?\d{1,3}[-\s]?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',
            r'\d{10}',
            r'\(\d{3}\)[-\s]?\d{3}[-\s]?\d{4}'
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return ''

    def _extract_skills(self, text):
        """Extract all skills from text"""
        found_skills = set()
        text_lower = text.lower()

        # Check all skill categories
        for category, skills in self.skill_categories.items():
            for skill in skills:
                # Use word boundary matching for accuracy
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.add(skill.title())

        return sorted(list(found_skills))

    def _extract_experience(self, text):
        """Extract work experience"""
        experience = []

        # Look for year patterns
        year_pattern = r'20\d{2}'
        matches = re.finditer(year_pattern, text)

        for match in matches:
            # Extract context around year
            start = max(0, match.start() - 200)
            end = min(len(text), match.end() + 200)
            context = text[start:end]

            # Look for job title indicators
            if any(keyword in context.lower() for keyword in 
                   ['engineer', 'developer', 'manager', 'analyst', 'designer', 'lead']):
                experience.append({
                    'year': match.group(0),
                    'context': context.strip()[:100]
                })

        return experience[:5]  # Return top 5

    def _extract_education(self, text):
        """Extract education information"""
        education = []

        edu_keywords = ['university', 'college', 'bachelor', 'master', 'phd', 
                        'degree', 'diploma', 'b.tech', 'm.tech', 'b.sc', 'm.sc',
                        'b.e', 'm.e', 'mba', 'bba']

        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in edu_keywords):
                education.append(line.strip())

        return education[:3]  # Return top 3

    def _extract_location(self, text):
        """Extract location"""
        # Common location patterns
        location_patterns = [
            r'\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z][a-z]+',  # City, State
            r'\b(?:Bangalore|Mumbai|Delhi|Hyderabad|Chennai|Pune|Kolkata)\b'
        ]

        for pattern in location_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)

        return ''

    def _extract_projects(self, text):
        """Extract project information"""
        projects = []

        # Look for project section
        project_pattern = r'(?:projects?|portfolio)\s*:?\s*([^\n]*(?:\n[^\n]*){0,3})'
        matches = re.finditer(project_pattern, text, re.IGNORECASE)

        for match in matches:
            project_text = match.group(1).strip()
            if len(project_text) > 20:
                projects.append(project_text[:200])

        return projects[:3]

    def _categorize_skills_manual(self, skills):
        """Categorize skills manually using predefined categories"""
        categorized = {category: [] for category in self.skill_categories.keys()}

        for skill in skills:
            skill_lower = skill.lower()
            for category, keywords in self.skill_categories.items():
                if skill_lower in keywords:
                    categorized[category].append(skill)
                    break

        # Remove empty categories
        return {k: v for k, v in categorized.items() if v}

    def _categorize_skills_with_ai(self, skills, full_text):
        """Use Groq AI to categorize skills more intelligently"""
        try:
            prompt = f"""Analyze these skills and categorize them into the following categories:
- Technical Skills (programming languages, frameworks)
- Tools & Platforms (DevOps, cloud, databases)
- Domain Knowledge (ML/AI, data science, web development)
- Soft Skills (communication, leadership, etc.)

Skills: {', '.join(skills)}

Respond ONLY with JSON in this exact format:
{{
    "technical_skills": ["skill1", "skill2"],
    "tools_platforms": ["skill1", "skill2"],
    "domain_knowledge": ["skill1", "skill2"],
    "soft_skills": ["skill1", "skill2"]
}}"""

            response = self.groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=500
            )

            result_text = response.choices[0].message.content

            # Extract JSON from response
            json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
            if json_match:
                categorized = json.loads(json_match.group())
                return categorized
            else:
                return self._categorize_skills_manual(skills)

        except Exception as e:
            print(f"AI categorization failed: {e}, using manual categorization")
            return self._categorize_skills_manual(skills)
